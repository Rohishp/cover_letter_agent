from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor

from cv.schema import BaseEntry, CVContent, Certification, Language, SkillGroup


FONT_NAME = "Calibri"
BODY_SIZE = Pt(10)
HEADING_SIZE = Pt(12)

DARK_CHARCOAL = RGBColor(0x33, 0x33, 0x33)
DARK_NAVY = RGBColor(0x1B, 0x2A, 0x4A)
ACCENT = RGBColor(0x2E, 0x74, 0xB5)
ACCENT_HEX = "2E74B5"

MARGIN = Inches(0.7)

HYPERLINK_RELATIONSHIP_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"

SECTION_HEADINGS = {
    "profile": "Profile",
    "skills": "Skills",
    "experience": "Experience",
    "projects": "Projects",
    "education": "Education",
    "certifications": "Certifications",
    "languages": "Languages",
    "coursework": "Coursework",
}
SECTION_HEADING_TEXTS = set(SECTION_HEADINGS.values())

FIRST_SECTION = "profile"
LAST_SECTIONS = ["certifications", "languages", "coursework"]
MIDDLE_SECTIONS = {"skills", "experience", "projects", "education"}
ALL_SECTIONS = {FIRST_SECTION, *MIDDLE_SECTIONS, *LAST_SECTIONS}

DEFAULT_SECTION_ORDER = [
    "profile",
    "skills",
    "experience",
    "projects",
    "education",
    "certifications",
    "languages",
    "coursework",
]


def validate_section_order(section_order: list[str]) -> None:
    if len(section_order) != len(set(section_order)):
        raise ValueError(f"section_order contains duplicates: {section_order}")

    if set(section_order) != ALL_SECTIONS:
        missing = ALL_SECTIONS - set(section_order)
        extra = set(section_order) - ALL_SECTIONS
        raise ValueError(
            f"section_order must contain exactly {sorted(ALL_SECTIONS)}; "
            f"missing={sorted(missing)} extra={sorted(extra)}"
        )

    if section_order[0] != FIRST_SECTION:
        raise ValueError(f"section_order must start with '{FIRST_SECTION}', got '{section_order[0]}'")

    if section_order[-3:] != LAST_SECTIONS:
        raise ValueError(f"section_order must end with {LAST_SECTIONS}, got {section_order[-3:]}")

    middle = section_order[1:-3]

    if set(middle) != MIDDLE_SECTIONS:
        raise ValueError(f"middle sections must be exactly {sorted(MIDDLE_SECTIONS)}, got {middle}")


def _add_run(
    paragraph,
    text: str,
    *,
    bold: bool = False,
    size=BODY_SIZE,
    color: RGBColor = DARK_CHARCOAL,
):
    run = paragraph.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = size
    run.font.bold = bold
    run.font.color.rgb = color
    return run


def _add_hyperlink(
    paragraph,
    url: str,
    text: str,
    *,
    size=BODY_SIZE,
    color: RGBColor = DARK_CHARCOAL,
) -> None:
    """
    Insert a real w:hyperlink relationship, styled like normal body text
    (no blue/underline) to stay inside the fixed body/heading/accent
    colour palette.
    """

    part = paragraph.part
    r_id = part.relate_to(url, HYPERLINK_RELATIONSHIP_TYPE, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT_NAME)
    r_fonts.set(qn("w:hAnsi"), FONT_NAME)
    r_pr.append(r_fonts)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size.pt * 2)))
    r_pr.append(sz)

    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), str(color))
    r_pr.append(color_el)

    run.append(r_pr)

    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _strip_url_scheme(url: str) -> str:
    display = url

    if display.startswith("https://"):
        display = display[len("https://"):]
    elif display.startswith("http://"):
        display = display[len("http://"):]

    if display.startswith("www."):
        display = display[len("www."):]

    return display


def _add_bottom_border(paragraph, color_hex: str, size_eighths_pt: int = 6) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size_eighths_pt))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color_hex)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _add_section_heading(doc: Document, section_name: str):
    text = SECTION_HEADINGS[section_name]

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.keep_with_next = True
    _add_run(para, text, bold=True, size=HEADING_SIZE, color=DARK_NAVY)
    _add_bottom_border(para, ACCENT_HEX)
    return para


def _add_bullet(doc: Document, text: str, space_after=Pt(2)):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = space_after
    para.paragraph_format.keep_together = True
    _add_run(para, text)
    return para


def _add_header_block(doc: Document, content: CVContent) -> None:
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_after = Pt(2)
    _add_run(name_para, content.full_name, bold=True, size=HEADING_SIZE, color=DARK_NAVY)

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Pt(2)
    _add_run(title_para, content.target_title)

    # Line 1: location, email, phone -- plain text, no links.
    line1_para = doc.add_paragraph()
    line1_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line1_para.paragraph_format.space_after = Pt(2)
    _add_run(line1_para, " | ".join([content.location, *content.contact_details]))

    # Line 2: linkedin/github -- real hyperlinks, scheme/www stripped for display.
    line2_para = doc.add_paragraph()
    line2_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line2_para.paragraph_format.space_after = Pt(10)

    for index, url in enumerate(content.contact_links):
        if index > 0:
            _add_run(line2_para, " | ")

        _add_hyperlink(line2_para, url, _strip_url_scheme(url))


def _add_entry(doc: Document, entry: BaseEntry) -> None:
    row1 = doc.add_paragraph()
    row1.paragraph_format.space_before = Pt(8)
    row1.paragraph_format.space_after = Pt(0)
    row1.paragraph_format.keep_with_next = True
    row1.paragraph_format.keep_together = True
    _add_run(row1, entry.position, bold=True)

    org_and_city = (
        f"{entry.organisation}, {entry.city}" if entry.city else entry.organisation
    )
    _add_run(row1, f" | {org_and_city} | {entry.dates}")

    if entry.tools:
        tools_para = doc.add_paragraph()
        tools_para.paragraph_format.space_before = Pt(0)
        tools_para.paragraph_format.space_after = Pt(0)
        tools_para.paragraph_format.keep_with_next = True
        tools_para.paragraph_format.keep_together = True
        _add_run(tools_para, f"Tools: {entry.tools}")

    if entry.topic:
        topic_para = doc.add_paragraph()
        topic_para.paragraph_format.space_before = Pt(0)
        topic_para.paragraph_format.space_after = Pt(2)
        topic_para.paragraph_format.keep_with_next = True
        topic_para.paragraph_format.keep_together = True
        _add_run(topic_para, entry.topic, bold=True)

    bullet_count = len(entry.bullets)

    for index, bullet_text in enumerate(entry.bullets):
        is_last = index == bullet_count - 1
        _add_bullet(doc, bullet_text, space_after=Pt(6) if is_last else Pt(2))


def _add_skill_group(doc: Document, group: SkillGroup) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(4)
    _add_run(para, f"{group.name}: ", bold=True)
    _add_run(para, ", ".join(group.items))


def _add_certification(doc: Document, cert: Certification) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(4)
    _add_run(para, cert.name, bold=True)

    detail = f" | {cert.issuer} | {cert.issued}"

    if cert.expires:
        detail += f" to {cert.expires}"

    _add_run(para, detail)


def _add_language(doc: Document, language: Language) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(4)
    _add_run(para, f"{language.language}: ", bold=True)
    _add_run(para, language.level)


def _set_default_style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = BODY_SIZE


def _render_profile(doc: Document, content: CVContent) -> None:
    if not content.profile:
        return

    _add_section_heading(doc, "profile")

    profile_para = doc.add_paragraph()
    profile_para.paragraph_format.space_after = Pt(6)
    _add_run(profile_para, content.profile)


def _render_skills(doc: Document, content: CVContent) -> None:
    if not content.skill_groups:
        return

    _add_section_heading(doc, "skills")

    for group in content.skill_groups:
        _add_skill_group(doc, group)


def _render_experience(doc: Document, content: CVContent) -> None:
    if not content.experience:
        return

    _add_section_heading(doc, "experience")

    for entry in content.experience:
        _add_entry(doc, entry)


def _render_projects(doc: Document, content: CVContent) -> None:
    if not content.projects:
        return

    _add_section_heading(doc, "projects")

    for entry in content.projects:
        _add_entry(doc, entry)


def _render_education(doc: Document, content: CVContent) -> None:
    if not content.education:
        return

    _add_section_heading(doc, "education")

    for entry in content.education:
        _add_entry(doc, entry)


def _render_certifications(doc: Document, content: CVContent) -> None:
    if not content.certifications:
        return

    _add_section_heading(doc, "certifications")

    for cert in content.certifications:
        _add_certification(doc, cert)


def _render_languages(doc: Document, content: CVContent) -> None:
    if not content.languages:
        return

    _add_section_heading(doc, "languages")

    for language in content.languages:
        _add_language(doc, language)


def _render_coursework(doc: Document, content: CVContent) -> None:
    if not content.coursework:
        return

    _add_section_heading(doc, "coursework")

    for item in content.coursework:
        _add_bullet(doc, item)


SECTION_RENDERERS = {
    "profile": _render_profile,
    "skills": _render_skills,
    "experience": _render_experience,
    "projects": _render_projects,
    "education": _render_education,
    "certifications": _render_certifications,
    "languages": _render_languages,
    "coursework": _render_coursework,
}


def render_cv(
    content: CVContent,
    out_path: Path,
    section_order: list[str] = DEFAULT_SECTION_ORDER,
) -> Path:
    """
    Render CVContent to a .docx file.

    Owns every formatting rule. Deterministic: same content in, same
    document out, every time.

    section_order controls render order: profile is always first,
    certifications/languages/coursework are always the last three (in
    that order), and skills/experience/projects/education may be
    reordered freely between them. Phase B will compute the middle
    order from the job description.
    """

    validate_section_order(section_order)

    doc = Document()

    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = MARGIN
    section.right_margin = MARGIN
    section.top_margin = MARGIN
    section.bottom_margin = MARGIN

    _set_default_style(doc)

    _add_header_block(doc, content)

    for section_name in section_order:
        SECTION_RENDERERS[section_name](doc, content)

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))

    return out_path
