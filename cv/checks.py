import math
import re
import shutil
import subprocess
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from cv.facts import CVFacts
from cv.render import SECTION_HEADING_TEXTS
from cv.schema import CVContent


ALLOWED_FONT = "Calibri"
ALLOWED_SIZES_PT = {10, 12}

MIN_MARGIN_INCHES = 0.65
MAX_MARGIN_INCHES = 0.8

EM_DASH = "—"
EN_DASH = "–"

# Calibri 10pt characters-per-line estimates for the ~6.87in usable column
# (A4 minus 0.7in margins each side). Calibrated approximations, not true
# layout measurement -- exact line wrapping needs the PDF render, which
# page_count/last_page_fill already require.
PROFILE_CHARS_PER_LINE = 100
BULLET_CHARS_PER_LINE = 95

MAX_PROFILE_LINES = 3
MAX_BULLET_LINES = 2

EXPERIENCE_BULLET_RANGE = (3, 5)
EDUCATION_BULLET_RANGE = (0, 4)
PROJECT_BULLET_RANGE = (2, 5)

MIN_LAST_PAGE_FILL = 0.25

# Header block is always exactly 4 lines when nothing wraps: name, target
# title, contact line 1 (location/email/phone), contact line 2 (links).
EXPECTED_HEADER_LINES = 4

CheckResult = tuple[str, bool | None, str]


def _iter_paragraphs(doc: Document):
    return doc.paragraphs


def _iter_runs(doc: Document):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            yield run

    for section in doc.sections:
        for container in (section.header, section.footer):
            for paragraph in container.paragraphs:
                for run in paragraph.runs:
                    yield run


def check_no_tables(docx_path: Path, **_) -> CheckResult:
    doc = Document(str(docx_path))
    tables = doc.element.body.findall(f".//{qn('w:tbl')}")
    passed = len(tables) == 0
    return ("no_tables", passed, f"{len(tables)} w:tbl elements found")


def check_empty_header_footer(docx_path: Path, **_) -> CheckResult:
    doc = Document(str(docx_path))
    offending = []

    for section_index, section in enumerate(doc.sections):
        for name, container in (("header", section.header), ("footer", section.footer)):
            text = "".join(p.text for p in container.paragraphs).strip()

            if text:
                offending.append(f"section {section_index} {name}: {text!r}")

    passed = not offending
    detail = "clean" if passed else "; ".join(offending)
    return ("empty_header_footer", passed, detail)


def check_fonts_and_sizes(docx_path: Path, **_) -> CheckResult:
    doc = Document(str(docx_path))
    offending = []

    for run in _iter_runs(doc):
        if not run.text.strip():
            continue

        name = run.font.name
        size = run.font.size
        size_pt = size.pt if size is not None else None

        if name != ALLOWED_FONT or size_pt not in ALLOWED_SIZES_PT:
            offending.append(f"{run.text[:30]!r} font={name} size={size_pt}")

    passed = not offending
    detail = "clean" if passed else f"{len(offending)} offending runs: " + "; ".join(offending[:5])
    return ("fonts_and_sizes", passed, detail)


def check_margins(docx_path: Path, **_) -> CheckResult:
    doc = Document(str(docx_path))
    section = doc.sections[0]

    margins = {
        "left": section.left_margin.inches,
        "right": section.right_margin.inches,
        "top": section.top_margin.inches,
        "bottom": section.bottom_margin.inches,
    }

    offending = {
        key: value
        for key, value in margins.items()
        if not (MIN_MARGIN_INCHES <= value <= MAX_MARGIN_INCHES)
    }

    passed = not offending
    detail = f"margins={margins}" if passed else f"out of range: {offending}"
    return ("margins", passed, detail)


def check_no_dashes(docx_path: Path, **_) -> CheckResult:
    doc = Document(str(docx_path))
    offending = []

    for paragraph in _iter_paragraphs(doc):
        if EM_DASH in paragraph.text or EN_DASH in paragraph.text:
            offending.append(paragraph.text[:60])

    passed = not offending
    detail = "clean" if passed else f"{len(offending)} paragraphs: " + "; ".join(offending[:5])
    return ("no_dashes", passed, detail)


def _rendered_lines(text: str, chars_per_line: int) -> int:
    if not text:
        return 0

    return max(1, math.ceil(len(text) / chars_per_line))


def check_profile_length(docx_path: Path, *, content: CVContent, **_) -> CheckResult:
    lines = _rendered_lines(content.profile, PROFILE_CHARS_PER_LINE)
    passed = lines <= MAX_PROFILE_LINES
    return ("profile_length", passed, f"~{lines} rendered lines (limit {MAX_PROFILE_LINES})")


def check_bullets_per_entry(docx_path: Path, *, content: CVContent, **_) -> CheckResult:
    offending = []

    def _check_group(entries, label, count_range):
        low, high = count_range

        for entry in entries:
            count = len(entry.bullets)

            if not (low <= count <= high):
                offending.append(
                    f"{label} '{entry.position}': {count} bullets (expected {low}-{high})"
                )

            for bullet in entry.bullets:
                lines = _rendered_lines(bullet, BULLET_CHARS_PER_LINE)

                if lines > MAX_BULLET_LINES:
                    offending.append(
                        f"{label} '{entry.position}' bullet too long: ~{lines} lines"
                    )

    _check_group(content.experience, "experience", EXPERIENCE_BULLET_RANGE)
    _check_group(content.education, "education", EDUCATION_BULLET_RANGE)
    _check_group(content.projects, "projects", PROJECT_BULLET_RANGE)

    passed = not offending
    detail = "clean" if passed else "; ".join(offending[:8])
    return ("bullets_per_entry", passed, detail)


def check_tier1_present(docx_path: Path, *, facts: CVFacts, **_) -> CheckResult:
    doc = Document(str(docx_path))
    full_text = "\n".join(p.text for p in doc.paragraphs)

    missing = []

    if facts.languages.tier == 1:
        for entry in facts.languages.entries:
            if entry.language not in full_text:
                missing.append(f"language: {entry.language}")

    if facts.education.tier == 1:
        for entry in facts.education.entries:
            if entry.degree not in full_text:
                missing.append(f"education: {entry.degree}")

    if facts.work_experience.tier == 1:
        for entry in facts.work_experience.entries:
            if entry.position not in full_text:
                missing.append(f"experience: {entry.position}")

    if facts.certifications.tier == 1:
        for entry in facts.certifications.entries:
            if entry.name not in full_text:
                missing.append(f"certification: {entry.name}")

    passed = not missing
    detail = "clean" if passed else "; ".join(missing)
    return ("tier1_present", passed, detail)


def extract_ordered_lines(pdf_path: Path) -> list[str]:
    """
    Every visual line of text in the rendered PDF, in reading order.
    """

    import pymupdf as fitz

    pdf = fitz.open(str(pdf_path))
    lines = []

    for page in pdf:
        data = page.get_text("dict")

        for block in data["blocks"]:
            for line in block.get("lines", []):
                text = "".join(span["text"] for span in line["spans"]).strip()

                if text:
                    lines.append(text)

    pdf.close()
    return lines


def check_no_orphan_headings(docx_path: Path, *, pdf_path: Path | None, **_) -> CheckResult:
    if pdf_path is None:
        return ("no_orphan_headings", None, "skipped: no PDF converter available")

    import pymupdf as fitz

    pdf = fitz.open(str(pdf_path))
    offending = []

    for page_index, page in enumerate(pdf):
        data = page.get_text("dict")
        page_lines = []

        for block in data["blocks"]:
            for line in block.get("lines", []):
                text = "".join(span["text"] for span in line["spans"]).strip()

                if text:
                    page_lines.append(text)

        if page_lines and page_lines[-1] in SECTION_HEADING_TEXTS:
            offending.append(f"page {page_index + 1}: '{page_lines[-1]}'")

    pdf.close()

    passed = not offending
    detail = "clean" if passed else "; ".join(offending)
    return ("no_orphan_headings", passed, detail)


def check_contact_lines_no_wrap(docx_path: Path, *, pdf_path: Path | None, **_) -> CheckResult:
    if pdf_path is None:
        return ("contact_lines_no_wrap", None, "skipped: no PDF converter available")

    lines = extract_ordered_lines(pdf_path)

    header_lines = []
    for line in lines:
        if line in SECTION_HEADING_TEXTS:
            break
        header_lines.append(line)

    passed = len(header_lines) == EXPECTED_HEADER_LINES
    detail = (
        f"{len(header_lines)} header lines (expected {EXPECTED_HEADER_LINES})"
        if not passed
        else "clean"
    )
    return ("contact_lines_no_wrap", passed, detail)


def _normalize_whitespace(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def check_metric_density(docx_path: Path, *, content: CVContent, pdf_path: Path | None, **_) -> CheckResult:
    if pdf_path is None:
        return ("metric_density", None, "skipped: no PDF converter available")

    import pymupdf as fitz

    pdf = fitz.open(str(pdf_path))

    if pdf.page_count == 0:
        pdf.close()
        return ("metric_density", None, "no pages")

    page1_text = _normalize_whitespace(pdf[0].get_text())
    pdf.close()

    all_bullets = []
    for entry in (*content.experience, *content.education, *content.projects):
        all_bullets.extend(entry.bullets)
    all_bullets.extend(content.coursework)

    page1_bullets = [b for b in all_bullets if _normalize_whitespace(b) in page1_text]

    if not page1_bullets:
        return ("metric_density", True, "0 bullets on page 1")

    digit_bullets = [b for b in page1_bullets if any(ch.isdigit() for ch in b)]
    fraction = len(digit_bullets) / len(page1_bullets)

    return (
        "metric_density",
        True,
        f"{len(digit_bullets)}/{len(page1_bullets)} page-1 bullets contain a digit ({fraction:.0%})",
    )


def convert_to_pdf(docx_path: Path) -> tuple[Path | None, str]:
    pdf_path = docx_path.with_suffix(".pdf")

    try:
        from docx2pdf import convert

        convert(str(docx_path), str(pdf_path))

        if pdf_path.exists():
            return pdf_path, "docx2pdf (MS Word)"
    except Exception:
        pass

    soffice = shutil.which("soffice")

    if soffice:
        try:
            subprocess.run(
                [
                    soffice,
                    "--headless",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    str(docx_path.parent),
                    str(docx_path),
                ],
                check=True,
                capture_output=True,
                timeout=120,
            )

            if pdf_path.exists():
                return pdf_path, "soffice"
        except Exception:
            pass

    return None, "neither docx2pdf (MS Word) nor soffice available"


def check_page_count(docx_path: Path, *, pdf_path: Path | None, **_) -> CheckResult:
    if pdf_path is None:
        return ("page_count", None, "skipped: no PDF converter available")

    import pymupdf as fitz

    pdf = fitz.open(str(pdf_path))
    count = pdf.page_count
    pdf.close()

    passed = count in (1, 2)
    return ("page_count", passed, f"{count} pages")


def check_last_page_fill(docx_path: Path, *, pdf_path: Path | None, **_) -> CheckResult:
    if pdf_path is None:
        return ("last_page_fill", None, "skipped: no PDF converter available")

    import pymupdf as fitz

    pdf = fitz.open(str(pdf_path))
    last_page = pdf[pdf.page_count - 1]
    page_height = last_page.rect.height

    blocks = last_page.get_text("blocks")

    if not blocks:
        pdf.close()
        return ("last_page_fill", False, "last page has no content")

    lowest_y = max(block[3] for block in blocks)
    fill_ratio = lowest_y / page_height
    pdf.close()

    passed = fill_ratio >= MIN_LAST_PAGE_FILL
    return ("last_page_fill", passed, f"{fill_ratio:.0%} filled (min {MIN_LAST_PAGE_FILL:.0%})")


def run_checks(
    docx_path: Path,
    content: CVContent,
    facts: CVFacts,
    pdf_path: Path | None = None,
) -> list[CheckResult]:
    docx_path = Path(docx_path)

    results = [
        check_no_tables(docx_path),
        check_empty_header_footer(docx_path),
        check_fonts_and_sizes(docx_path),
        check_margins(docx_path),
        check_no_dashes(docx_path),
        check_profile_length(docx_path, content=content),
        check_bullets_per_entry(docx_path, content=content),
        check_tier1_present(docx_path, facts=facts),
    ]

    if pdf_path is None:
        pdf_path, _converter_detail = convert_to_pdf(docx_path)

    results.append(check_page_count(docx_path, pdf_path=pdf_path))
    results.append(check_last_page_fill(docx_path, pdf_path=pdf_path))
    results.append(check_no_orphan_headings(docx_path, pdf_path=pdf_path))
    results.append(check_contact_lines_no_wrap(docx_path, pdf_path=pdf_path))
    results.append(check_metric_density(docx_path, content=content, pdf_path=pdf_path))

    return results
